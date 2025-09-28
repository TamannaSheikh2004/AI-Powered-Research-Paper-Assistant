import os
import io
import re
import requests
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
import urllib.parse

load_dotenv()
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "").strip()

def generate_text_grok(prompt: str, max_tokens: int = 2000) -> str:
    if not OPENROUTER_API_KEY:
        return "[Error: API key missing]"
    
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8501"
    }
    
    body = {
        "model": "x-ai/grok-4-fast:free",
        "messages": [
            {"role": "system", "content": "You are an expert academic researcher. Generate detailed, high-quality research content."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens
    }
    
    try:
        response = requests.post(url, headers=headers, json=body, timeout=120)
        response.raise_for_status()
        data = response.json()
        content = data["choices"][0]["message"]["content"].strip()
        return content if content else "[Error: Empty response]"
    except Exception as e:
        return f"[Error: {str(e)}]"

def generate_image_pollinations(prompt: str, width: int = 1024, height: int = 768) -> bytes:
    try:
        clean_prompt = re.sub(r'[^\w\s\-.,]', '', prompt)
        enhanced_prompt = f"{clean_prompt}, professional, high quality, detailed, research illustration"
        encoded_prompt = urllib.parse.quote(enhanced_prompt)
        
        url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width={width}&height={height}&model=flux&enhance=true&nologo=true"
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        
        response = requests.get(url, headers=headers, timeout=120, stream=True)
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '')
        if not content_type.startswith('image/'):
            alt_url = f"https://pollinations.ai/p/{encoded_prompt}?width={width}&height={height}"
            response = requests.get(alt_url, headers=headers, timeout=120, stream=True)
            response.raise_for_status()
        
        image_data = response.content
        
        if len(image_data) < 1000:
            raise RuntimeError("Generated image too small")
        
        try:
            img = Image.open(io.BytesIO(image_data))
            img.verify()
        except Exception:
            raise RuntimeError("Invalid image data received")
        
        return image_data
        
    except requests.exceptions.Timeout:
        raise RuntimeError("Image generation timed out - server busy")
    except requests.exceptions.RequestException as e:
        raise RuntimeError(f"Network error during image generation: {str(e)}")
    except Exception as e:
        raise RuntimeError(f"Image generation failed: {str(e)}")

def test_pollinations_api():
    try:
        test_image = generate_image_pollinations("test image", 256, 256)
        return True
    except:
        return False

def create_placeholder_image(text: str, width=800, height=450) -> io.BytesIO:
    img = Image.new("RGB", (width, height), color=(240, 248, 255))
    draw = ImageDraw.Draw(img)
    
    font = ImageFont.load_default()
    draw.rectangle([10, 10, width-10, height-10], outline=(70, 130, 180), width=3)
    draw.text((50, 30), "Research Illustration", font=font, fill=(25, 25, 112))
    
    words = text.split()
    lines = []
    current_line = ""
    
    for word in words:
        test_line = current_line + " " + word if current_line else word
        if len(test_line) <= 60:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    
    if current_line:
        lines.append(current_line)
    
    y_offset = 80
    for line in lines[:8]:
        draw.text((30, y_offset), line, font=font, fill=(60, 60, 60))
        y_offset += 30
    
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer

def generate_graph(description: str, graph_type: str = "bar") -> io.BytesIO:
    plt.style.use('default')
    fig, ax = plt.subplots(figsize=(10, 6))
    
    if graph_type == "bar":
        categories = ["Method A", "Method B", "Method C", "Proposed"]
        values = [0.75, 0.82, 0.88, 0.94]
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99']
        bars = ax.bar(categories, values, color=colors)
        ax.set_ylabel('Accuracy')
        ax.set_title('Performance Comparison')
        
        for bar, val in zip(bars, values):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                   f'{val:.2f}', ha='center', va='bottom')
    
    elif graph_type == "line":
        x = np.arange(2019, 2025)
        y = [65, 72, 78, 85, 89, 92]
        ax.plot(x, y, marker='o', linewidth=3, markersize=8, color='#2E86AB')
        ax.set_xlabel('Year')
        ax.set_ylabel('Performance (%)')
        ax.set_title('Performance Over Time')
        ax.grid(True, alpha=0.3)
    
    elif graph_type == "pie":
        sizes = [35, 25, 20, 15, 5]
        labels = ['Type A', 'Type B', 'Type C', 'Type D', 'Others']
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
        ax.set_title('Data Distribution')
    
    plt.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="PNG", dpi=300, bbox_inches='tight')
    plt.close(fig)
    buffer.seek(0)
    return buffer

def process_template(uploaded_template, sections):
    if uploaded_template:
        doc = Document(uploaded_template)
        new_doc = Document()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            section_found = None
            for section_name in sections.keys():
                if (text.lower() == section_name.lower() or 
                    section_name.lower() in text.lower() or 
                    text.lower() in section_name.lower()):
                    section_found = section_name
                    break
            
            if section_found and sections[section_found]:
                new_doc.add_heading(text, level=1)
                content = sections[section_found]
                if not content.startswith("[Error"):
                    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                    for para_text in paragraphs:
                        if para_text:
                            new_para = new_doc.add_paragraph(para_text)
                            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif text:
                if "title" in para.style.name.lower():
                    new_doc.add_heading(text, level=0)
                else:
                    new_doc.add_paragraph(text)
        
        return new_doc
    else:
        return None

st.set_page_config(page_title="AI Research Assistant", layout="wide")

st.title("AI Research Assistant")
st.info("Generate comprehensive research papers with AI-powered content and visuals")

with st.sidebar:
    st.header("Status")
    if OPENROUTER_API_KEY:
        st.success("OpenRouter API Ready")
    else:
        st.error("API Key Missing")
    
    st.write("Testing Image Generation...")
    if test_pollinations_api():
        st.success("Pollinations AI Ready")
    else:
        st.warning("Pollinations API Issues")

topic = st.text_input("Research Topic", placeholder="Enter your research topic")

st.subheader("Document Template")
template_option = st.radio("Template", ["Default Format", "Upload Template"])

uploaded_template = None
if template_option == "Upload Template":
    uploaded_template = st.file_uploader("Upload DOCX Template", type=["docx"])

st.subheader("Content Structure")
default_sections = ["Abstract", "Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion", "References"]
sections_input = st.text_area("Sections (one per line)", value="\n".join(default_sections), height=150)
user_sections = [s.strip() for s in sections_input.split('\n') if s.strip()]

st.subheader("Visual Content")
col1, col2, col3 = st.columns(3)

with col1:
    st.write("Research Figures")
    fig_option = st.selectbox("Figures", ["None", "Generate", "Upload"])
    fig_prompt = ""
    uploaded_figures = []
    if fig_option == "Generate":
        fig_prompt = st.text_input("Describe figure", placeholder="system architecture, process flow")
    elif fig_option == "Upload":
        uploaded_figures = st.file_uploader("Upload figures", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

with col2:
    st.write("Data Graphs")
    graph_option = st.selectbox("Graphs", ["None", "Generate", "Upload"])
    graph_prompt = ""
    graph_type = "bar"
    uploaded_graphs = []
    if graph_option == "Generate":
        graph_prompt = st.text_input("Describe data", placeholder="performance metrics, comparison")
        graph_type = st.selectbox("Graph Type", ["bar", "line", "pie"])
    elif graph_option == "Upload":
        uploaded_graphs = st.file_uploader("Upload graphs", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

with col3:
    st.write("Cover Image")
    cover_option = st.selectbox("Cover", ["None", "Generate", "Upload"])
    cover_prompt = ""
    uploaded_cover = None
    if cover_option == "Generate":
        cover_prompt = st.text_input("Describe cover", placeholder="research theme visualization")
    elif cover_option == "Upload":
        uploaded_cover = st.file_uploader("Upload cover", type=["png", "jpg", "jpeg"])

if st.button("Generate Research Paper", type="primary"):
    if not topic.strip():
        st.error("Please enter a research topic!")
    elif not OPENROUTER_API_KEY:
        st.error("API key required in .env file!")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            status_text.text("Generating content...")
            progress_bar.progress(20)
            
            section_prompts = {
                "Abstract": f"Write a comprehensive abstract for '{topic}'. Include background, objectives, methodology, key results, and conclusions. Length: 200-250 words.",
                "Introduction": f"Write a detailed introduction for '{topic}' research. Cover background, current state, challenges, research gap, objectives. Length: 400-500 words.",
                "Literature Review": f"Write a literature review for '{topic}'. Discuss existing work, methodologies, gaps. Include citations [1], [2]. Length: 350-400 words.",
                "Methodology": f"Write methodology for '{topic}' research. Describe approach, design, implementation, evaluation. Length: 300-400 words.",
                "Results": f"Write results for '{topic}'. Present findings, metrics, analysis. Length: 300-350 words.",
                "Discussion": f"Write discussion for '{topic}'. Analyze implications, limitations, comparisons. Length: 250-300 words.",
                "Conclusion": f"Write conclusion for '{topic}'. Summarize contributions, findings, future work. Length: 200-250 words.",
                "References": f"Generate 10 realistic references for '{topic}' in IEEE format."
            }
            
            paper_sections = {}
            successful_sections = 0
            
            for i, section in enumerate(user_sections):
                status_text.text(f"Generating: {section}...")
                prompt = section_prompts.get(section, f"Write a detailed {section} section for '{topic}' research.")
                content = generate_text_grok(prompt, max_tokens=2000)
                paper_sections[section] = content
                
                if not content.startswith("[Error") and len(content) > 100:
                    successful_sections += 1
                
                progress_bar.progress(20 + (i * 40) // len(user_sections))
            
            if successful_sections == 0:
                st.error("Failed to generate content. Check your API key and try again.")
                st.info(f"API Key Status: {'Set' if OPENROUTER_API_KEY else 'Missing'}")
                st.stop()
            
            status_text.text("Creating visuals...")
            progress_bar.progress(70)
            
            generated_images = []
            
            if fig_option == "Generate" and fig_prompt:
                status_text.text("Generating research figure...")
                try:
                    img_bytes = generate_image_pollinations(fig_prompt)
                    generated_images.append(("Research Figure", io.BytesIO(img_bytes)))
                    st.success("Figure generated successfully!")
                except Exception as e:
                    st.warning(f"Image generation failed: {str(e)}, using placeholder")
                    generated_images.append(("Research Figure", create_placeholder_image(fig_prompt)))
            
            if fig_option == "Upload" and uploaded_figures:
                for img in uploaded_figures:
                    generated_images.append(("Uploaded Figure", img))
            
            if graph_option == "Generate" and graph_prompt:
                status_text.text("Creating data graph...")
                graph_buffer = generate_graph(graph_prompt, graph_type)
                generated_images.append(("Data Graph", graph_buffer))
                st.success("Graph created successfully!")
            
            if graph_option == "Upload" and uploaded_graphs:
                for graph in uploaded_graphs:
                    generated_images.append(("Uploaded Graph", graph))
            
            if cover_option == "Generate" and cover_prompt:
                status_text.text("Generating cover image...")
                try:
                    cover_bytes = generate_image_pollinations(cover_prompt + " academic research style")
                    generated_images.append(("Cover Image", io.BytesIO(cover_bytes)))
                    st.success("Cover image generated successfully!")
                except Exception as e:
                    st.warning(f"Cover generation failed: {str(e)}, using placeholder")
                    generated_images.append(("Cover Image", create_placeholder_image(cover_prompt)))
            
            if cover_option == "Upload" and uploaded_cover:
                generated_images.append(("Cover Image", uploaded_cover))
            
            status_text.text("Creating document...")
            progress_bar.progress(85)
            
            if uploaded_template:
                doc = process_template(uploaded_template, paper_sections)
            else:
                doc = Document()
                title_para = doc.add_heading(topic, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                author_para = doc.add_paragraph("Author Name")
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_para.runs[0].italic = True
                
                date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
                
                for section, content in paper_sections.items():
                    if not content.startswith("[Error") and content.strip():
                        doc.add_heading(section, level=1)
                        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                        for para_text in paragraphs:
                            if para_text:
                                para = doc.add_paragraph(para_text)
                                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if generated_images:
                if not uploaded_template:
                    doc.add_heading("Figures and Visualizations", level=1)
                
                for img_name, img_buffer in generated_images:
                    try:
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(6))
                        caption = doc.add_paragraph(f"Figure: {img_name}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].italic = True
                        doc.add_paragraph()
                    except Exception:
                        pass
            
            status_text.text("Finalizing...")
            progress_bar.progress(95)
            
            safe_topic = re.sub(r'[^\w\s-]', '', topic).strip()[:25]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            filename = f"Research_Paper_{safe_topic.replace(' ', '_')}_{timestamp}.docx"
            
            doc.save(filename)
            
            progress_bar.progress(100)
            status_text.text("Complete!")
            
            with open(filename, "rb") as file:
                st.download_button(
                    label="Download Research Paper",
                    data=file,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.success("Research paper generated successfully!")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Sections Generated", successful_sections)
            with col2:
                st.metric("Visuals Added", len(generated_images))
            with col3:
                total_words = sum(len(content.split()) for content in paper_sections.values() 
                                if not content.startswith("[Error"))
                st.metric("Total Words", total_words)
            
            with st.expander("Content Preview"):
                for section, content in paper_sections.items():
                    if not content.startswith("[Error") and len(content) > 50:
                        st.subheader(section)
                        preview = content[:400] + "..." if len(content) > 400 else content
                        st.write(preview)
                        st.markdown("---")
            
        except Exception as e:
            st.error(f"Generation failed: {str(e)}")

st.markdown("---")
st.markdown("*AI Research Assistant - Made by Tamanna Sheikh*")